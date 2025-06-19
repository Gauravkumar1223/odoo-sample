import base64
import logging
import os
import uuid
from io import BytesIO

import pypandoc
from docx import Document
from docx.shared import Inches
from xhtml2pdf import pisa

from . import prompts
from .azur_llm_pool import AzurePoolLLM
from .azure_llm import AzureLLM

_logger = logging.getLogger(__name__)
logging.getLogger("fontTools").setLevel(logging.WARNING)
logging.getLogger("weasyprint").setLevel(logging.WARNING)


class VoiceRecordAiActions:
    wrrrit_llm = AzurePoolLLM(stream=False)


    @staticmethod
    def format_messages(messages):
        formatted_messages = ""
        for message in messages:
            role = message["role"]
            content = message["content"]
            if role == "system":
                formatted_messages += f"system\n\n{content}\n"
            elif role == "user":
                formatted_messages += f"user\n{content}\nInput\n"
            else:  # role == "assistant"
                formatted_messages += f"assistant\n{content}\n"
        return formatted_messages

    @staticmethod
    def printresponse(response):
        completion_text = ""
        # iterate through the stream of events and print it
        for event in response:
            event_text = event["choices"][0]["delta"]
            if "content" in event_text:
                chunk = event_text.content
                completion_text += chunk
                print(f"{chunk}", end="", flush=True)
        print("", flush=True)
        # remember context

        return completion_text

    @staticmethod
    def execute_openai_chat_completion(
            messages, temperature=0, max_tokens=8000, model="gpt-3.5-turbo-16k"
    ):
        import time, json, sys

        start_time = time.time()

        messages_json = json.dumps(messages, ensure_ascii=False)

        # Calculate the size of the JSON string in bytes
        message_size_bytes = round(sys.getsizeof(messages_json.encode("utf-8")) / 3)

        # Deduce max_tokens based on desired maximum token limit (e.g., 4096 - message_size_bytes)
        tokens = max_tokens - message_size_bytes
        wrrrit_llm = AzureLLM(stream=False)
        # Execute the chat completion
        response = wrrrit_llm.call_llm(
            messages=messages,
            temperature=temperature,
            max_tokens=tokens,
        )

        # Record the end time
        end_time = time.time()

        # Cost per 1000 tokens is $0.003 for this model

        # Calculate the time difference and log it
        time_taken = end_time - start_time
        _logger.info(f"Completed LLM Call. Time taken: {time_taken:.2f} seconds.")

        # Return the assistant's response
        result = response

        return result

    def build_system_prompt(
            self, prompt_template, report_template, text, locale="english"
    ):
        try:
            prompt_prefix = prompt_template.prompt_prefix or ""
            prompt_sections = prompt_template.prompt_sections or ""
            prompt_postfix = prompt_template.prompt_postfix or ""

            template_definition = report_template.template_definition or ""

            # Checking if we should include the current date declaration
            date_declaration = ""
            if prompt_template.insert_date:
                from datetime import datetime

                date_declaration = f"Note: Today  Document Date is: {datetime.now().strftime('%Y-%m-%d')}\n"

            # Checking if we should perform locale translation
            locale_declaration = ""
            if prompt_template.insert_locale:
                locale_declaration = f"Important Note: Translate non-{locale} text and data to {locale} everywhere.\n"

            # Building the system prompt
            system_prompt = (
                "Generate the response using strictly the following requirements:\n"
                f"{locale_declaration}"
                f"{date_declaration}"
                f"{prompt_prefix}\n"
                f"Sections: \n\n{prompt_sections}\n"
                f"{template_definition}\n"
                f"{prompt_postfix}\n"
            )

            # Preparing the user prompt
            user_prompt = f"Text: {text}"

            # Creating the messages array for ChatGPT
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ]
            # Generating report
            # Calling ChatGPT

            _logger.info("Calling LLM for report generation.")
            response = self.wrrrit_llm.call_llm(messages, max_tokens=8192)

            return response
            # return "Mucking GPT"
        except Exception as e:
            _logger.error(
                f"An error occurred while building the system prompt or calling ChatGPT: {e}"
            )
            return None  # Returning None or handle the error as needed

    def build_individual_section_report(
            self,
            section_title,
            section_content,
            system_prompt,
            user_prompt,
            locale="english",
    ):
        try:
            # Construct the system prompt for the specific section
            system_prompt = f"{system_prompt}\n\n" f"Sections: {section_title}\n\n"

            # Prepare the user prompt for the specific section
            user_prompt = f"Text: {section_content}"

            # Create the messages array for ChatGPT
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ]

            # Calling ChatGPT for the specific section
            _logger.info(f"Calling ChatGPT for section: {section_title}")
            response = self.wrrrit_llm.call_llm(messages, max_tokens=4000)

            return response
        except Exception as e:
            _logger.error(
                f"An error occurred while processing section '{section_title}': {e}"
            )
            return None  # Returning None or handle the error as needed

    @staticmethod
    def text_to_pdf(text):
        import pdfkit
        import base64

        # Convert Markdown text to HTML
        html_content = text

        # Set options for proper encoding
        options = {
            "encoding": "UTF-8",
        }

        # Convert HTML to PDF using pdfkit
        pdf_bytes = pdfkit.from_string(html_content, False, options=options)

        # Convert bytes to base64 encoded bytes for storing in Odoo Binary field
        pdf_b64 = base64.b64encode(pdf_bytes)

        return pdf_b64

    @staticmethod
    def text_to_pdf_formatted(html_text, header, footer, logo):
        # CSS styles to include
        # CSS styles to include
        css_styles = """
            /* Container for all sections */
            
           
            
                /* Logo styles */
        .logo-container {
            text-align: right; /* right align the logo */
            margin-top: 2px; /* Margin at the top */
        }
        .header_logo {
            width: 140px; /* Set a fixed width for the logo */
            height: auto; /* Maintain aspect ratio */
        }
            .sections-container {
                margin: 10px;
                padding: 10px;
                
                border: 0px solid #ddd;
                border-radius: 4px;
            }

            /* Header for each section */
            .section-header {
                font-size: 18px;
                color: #333;
                margin-bottom: 10px;
                border-bottom: 0px solid #eaeaea;
                padding-bottom: 5px;
            }

            /* Content of each section */
            .section-content {
                font-size: 14px;
                color: #666;
                margin-bottom: 10px;
                padding: 5px;
            }

            /* Separator between sections */
            .section-separator {
                border-top: 0px solid #eaeaea;
                margin-top: 10px;
                margin-bottom: 10px;
            }

            /* Bulleted List */
            .section-bullet-list {
                list-style-type: disc;
                margin-left: 20px;
                color: #333;
            }

            /* Text Paragraphs */
            .section-text {
                font-size: 12px;
                color: #333;
                text-align: justify;
                line-height: 1.5;
            }

            /* Citations */
            .section-citation {
                font-style: italic;
                color: #888;
                padding-left: 15px;
                border-left: 0px solid #ccc;
                margin: 10px 0;
            }

            /* Footer */
            .section-footer {
                font-size: 12px;
                text-align: center;
                margin-top: 10px;
                color: #666;
            }
            /* Base styling for the report */
                .rep_container {
                    font-family: Arial, sans-serif;
                    color: #333;
                    line-height: 1.6;
                    padding: 20px;
                    margin: auto;
                    max-width: 800px;
                }
                
                /* Title */
                .rep_title {
                    font-size: 24px;
                    text-align: center;
                    color: #0056b3;
                    margin-bottom: 20px;
                    font-weight: bold;
                }
                
                /* Headers */
                .rep_header {
                    font-size: 20px;
                    color: #004085;
                    margin-top: 20px;
                    margin-bottom: 10px;
                    font-weight: bold;
                }
                
                /* Section */
                .rep_section {
                    margin-bottom: 15px;
                }
                
                /* Content */
                .rep_content {
                    text-align: justify;
                    margin-bottom: 10px;
                }
                
                /* Bullet Points */
                .rep_bullets {
                    list-style-type: disc;
                    margin-left: 20px;
                }
                
                .rep_bullets li {
                    margin-bottom: 5px;
                }
                
                /* Signature */
                .rep_signature {
                    text-align: right;
                    margin-top: 30px;
                    font-style: italic;
                }
                
                /* Signature Line */
                .rep_signature-line {
                    border-top: 1px solid #000;
                    width: 200px;
                    margin-top: 5px;
                    margin-bottom: 5px;
                    text-align: center;
                    margin-left: auto; /* Aligns the line to the right */
                }
                
                /* Signature Name */
                .rep_signature-name {
                    font-size: 16px;
                }
                
                /* Footer */
                .rep_footer {
                    text-align: center;
                    margin-top: 30px;
                    font-size: 12px;
                    color: #6c757d;
                }
                
                /* Report Table */
                .rep_table {
                    width: 100%;
                    border-collapse: collapse;
                    margin-top: 20px;
                }
                
                .rep_table th,
                .rep_table td {
                    border: 1px solid #dee2e6;
                    padding: 8px;
                    text-align: left;
                }
                
                .rep_table th {
                    background-color: #f8f9fa;
                }
                
                .rep_table tr:nth-child(even) {
                    background-color: #f2f2f2;
                }

        """

        # Rest of the method remains the same

        # Set default values if parameters are not set
        html_text = html_text or ""
        header = header or ""
        footer = footer or ""

        try:
            logo_img_tag = ""
            if logo:
                encoded_logo = base64.b64decode(logo)
                logo_path = "logo.png"
                with open(logo_path, "wb") as f:
                    f.write(encoded_logo)
                logo_img_tag = f'<img src="{logo_path}" class="header_logo"/>'

            # Construct the full HTML with header, footer, content, and CSS
            full_html = f"""
            <html>
                <head>
                    <style>
                        {css_styles}
                    </style>
                </head>
                <body>
                  <div class="logo-container">
                    {logo_img_tag}
                </div>
                    <div class="sections-container">
                        <div id="pdf-header">
                       
                            <div class="section-header">
                                {header}
                            </div>
                        </div>
                        <div class="section-content">
                            {html_text}
                        </div>
                        <div id="pdf-footer">
                            <div class="section-footer">
                                {footer}
                            </div>
                        </div>
                    </div>
                </body>
            </html>
            """

            # Convert HTML to PDF
            pdf = BytesIO()
            pisa_status = pisa.CreatePDF(BytesIO(full_html.encode("UTF-8")), dest=pdf)

            # Check for errors
            if pisa_status.err:
                _logger.error("Failed to generate PDF")
                raise Exception("PDF generation error")

            pdf_bytes = pdf.getvalue()
            pdf_b64 = base64.b64encode(pdf_bytes)

            # Cleanup
            if logo:
                os.remove("logo.png")

            return pdf_b64
        except Exception as e:
            _logger.error(f"Failed to generate PDF: {str(e)}")
            if logo:
                os.remove("logo.png")
            return None

    @staticmethod
    def html_to_docx_formatted(html_text, header, footer, logo):
        try:
            # Convert HTML to Markdown
            # md_text = markdown.markdown(html_text)
            md_text = html_text
            # Add Header and Footer to the Markdown text
            header = header if header and len(header) > 1 else ""
            md_text = md_text if md_text and len(md_text) > 1 else ""
            footer = footer if footer and len(footer) > 1 else ""

            md_text = f"# {header}\n\n{md_text}\n\n{footer}"

            # Convert Markdown to DOCX using pypandoc
            doc_path = f"document_{uuid.uuid4()}.docx"
            pypandoc.convert_text(md_text, "docx", format="html", outputfile=doc_path)

            if logo:
                # Generate a unique file name for the logo
                logo_filename = f"logo_{uuid.uuid4()}.png"
                logo_path = os.path.join(os.getcwd(), logo_filename)

                with open(logo_path, "wb") as logo_file:
                    logo_file.write(base64.b64decode(logo))

                # Open the DOCX file and add the logo
                doc = Document(doc_path)
                doc.paragraphs[
                    0
                ].clear()  # Clear the first paragraph which had the 'False' statement
                run = doc.paragraphs[0].add_run()
                run.add_picture(
                    logo_path, width=Inches(1.5)
                )  # Adjust width as necessary
                doc.save(doc_path)
                os.remove(logo_path)  # Delete the logo file if it was created

            # Optionally: Convert the DOCX file to base64 if needed
            with open(doc_path, "rb") as doc_file:
                doc_b64 = base64.b64encode(doc_file.read())

            # Optional: remove the doc file after converting it to base64
            os.remove(doc_path)

            return doc_b64

        except Exception as e:
            logging.error(f"Failed to generate Word Document: {str(e)}")
            return None

    def action_generate_report(self, records):
        import datetime

        for record in records:
            _logger.info(
                f"Starting data extraction for record ID {record.id} using LLM Enterprise Servers"
                f"."
            )

            if not record.transcription_data or not record.record_locale:
                _logger.error(
                    f"Null value encountered in record ID {record.id}. Skipping this record."
                )
                continue
            input_data = record.transcription_data

            try:
                footer = (
                    record.report_template[0].footer if record.report_template else ""
                )
                header = (
                    record.report_template[0].header if record.report_template else ""
                )
                logo = record.report_template[0].logo if record.report_template else ""
                record.generated_report = self.build_system_prompt(
                    record.report_template[0].prompt_template_id[0],
                    record.report_template[0],
                    input_data,
                    record.record_locale,
                )

                _logger.info("Completed report generation for record ID %s.", record.id)

                pdf_data = VoiceRecordAiActions.text_to_pdf_formatted(
                    record.generated_report, footer=footer, header=header, logo=logo
                )
                logo = record.report_template[0].logo if record.report_template else ""
                docx_data = VoiceRecordAiActions.html_to_docx_formatted(
                    record.generated_report, footer=footer, header=header, logo=logo
                )

                record.generated_file = pdf_data
                record.generated_docx = docx_data
                formatted_date = datetime.datetime.now().strftime("%y%m%d")
                sequence_number = (
                        record.env["ir.sequence"].next_by_code(
                            "wrrrit.ai.voice_record.pdf.sequence"
                        )
                        or "0001"
                )
                pdf_filename = f"{record.name}-{formatted_date}-{sequence_number}.pdf"
                docx_filename = f"{record.name}-{formatted_date}-{sequence_number}.docx"
                record.write(
                    {
                        "generated_file_name": pdf_filename,
                        "generated_file": pdf_data,
                        "generated_docx_name": docx_filename,
                        "generated_docx": docx_data,
                    }
                )
                report_name = f"{record.name}-{formatted_date}-{sequence_number}-{record.record_locale}"
                record.env["wrrrit.ai.voice_record.pdf"].create(
                    {
                        "name": report_name,
                        "pdf_report_data": pdf_data,
                        "voice_record_id": record.id,
                    }
                )
            except Exception as e:
                _logger.error(
                    f"An error occurred during report generation for record ID {record.id}: {e}"
                )

    @staticmethod
    def get_translation_from_chatgpt3(text, locale="english"):
        system = prompts.translate_text_prompt(locale)
        user = prompts.user_translation_prompt(text)
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ]

        # Translation
        return VoiceRecordAiActions.wrrrit_llm.call_llm(messages, max_tokens=8000)

    def action_translate_transcript(self, record, locale="english"):
        for record in record:
            record.translated_data = self.get_translation_from_chatgpt3(
                record.transcription_data, record.record_locale
            )

    def action_drugs_diseases(self, record):
        for record in record:
            system = prompts.get_medicament_prompt(record.record_locale)
            user = prompts.user_html_prompt(record.transcription_data)
            messages = [
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ]
            # _logger.info("system message: %s", system)
            # _logger.info("user message: %s", user)

            response = self.wrrrit_llm.call_llm(messages, max_tokens=5000)
            record.drugs_diseases = response
