import base64
import datetime
import json
import logging
import os
import sys
import uuid

import pypandoc
from docx import Document
from docx.shared import Inches
from weasyprint import HTML

from . import prompts
from .llm import WrrritLLM

_logger = logging.getLogger(__name__)
logging.getLogger("fontTools").setLevel(logging.WARNING)
logging.getLogger("weasyprint").setLevel(logging.WARNING)


class VoiceRecordAiActions:
    DEFAULT_SYSTEM_PROMPT = (
        "You are a helpful, respectful, and honest assistant. "
        "Always answer as helpfully as possible, while being safe. "
        "Your answers should not include any harmful, unethical, racist, sexist, toxic, dangerous, or illegal content. "
        "Please ensure that your responses are socially unbiased and positive in nature."
        "\n\n"
        "If a question does not make any sense, or is not factually coherent, explain why instead of answering something not correct. "
        "If you don't know the answer to a question, please don't share false information."
    )

    @staticmethod
    def format_messages(messages):
        formatted_messages = ""
        for message in messages:
            role = message["role"]
            content = message["content"]
            formatted_messages += f"{role}\n{content}\n"
            if role == "user":
                formatted_messages += "Input\n"
        return formatted_messages

    @staticmethod
    def execute_openai_chat_completion(
            messages, temperature=0, max_tokens=8000, model="gpt-3.5-turbo-16k"
    ):
        OPENAI_LLM = True
        messages_json = json.dumps(messages, ensure_ascii=False)
        message_size_bytes = round(sys.getsizeof(messages_json.encode("utf-8")) / 3)
        tokens = max_tokens - message_size_bytes

        wrrrit_llm = WrrritLLM()

        return wrrrit_llm.call_llm(messages, max_tokens=8000)

    @staticmethod
    def get_information_from_chatgpt3(text, locale="english", sections="Not defined"):
        system = prompts.system_global_report(locale=locale, sections=sections)
        user = prompts.user_html_prompt(text)

        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ]
        wrrrit_llm = WrrritLLM()

        return wrrrit_llm.call_llm(messages, max_tokens=8000)

    @staticmethod
    def text_to_pdf_formatted(html_text, header, footer, logo):
        try:
            logo_img_tag = ""
            if logo:
                encoded_logo = base64.b64decode(logo)
                logo_path = "logo.png"
                with open(logo_path, "wb") as f:
                    f.write(encoded_logo)
                logo_img_tag = f'<img src="{logo_path}" class="logo"/>'

            full_html = (
                "<html>"
                "<head>"
                "<style>"
                "body {font-family: Arial, sans-serif; margin: 0; padding: 0;}"
                ".header, .footer {padding: 10px;}"
                ".logo {float: right; width: 100px; margin-right: 10px;}"
                ".header-content, .footer-content {overflow: auto;}"
                ".content {padding: 25px;}"
                "</style>"
                "</head>"
                "<body>"
                f"<div class='header'>{logo_img_tag}<div class='header-content'>{header}</div></div>"
                f"<div class='content'>{html_text}</div>"
                f"<div class='footer'><div class='footer-content'>{footer}</div></div>"
                "</body>"
                "</html>"
            )

            html = HTML(string=full_html, base_url="./", encoding="utf-8")
            pdf_bytes = html.write_pdf()
            pdf_b64 = base64.b64encode(pdf_bytes)

            if logo:
                os.remove("logo.png")
            return pdf_b64
        except Exception as e:
            _logger.error(f"Failed to generate PDF: {str(e)}")
            if logo:
                os.remove("logo.png")
            return None

    def action_generate_report(self, records):
        for record in records:
            try:
                sections = (
                    record.report_template[0].template_definition
                    if record.report_template
                    else (
                        "1. Patient Info\n"
                        "2. Medical History\n"
                        "3. Diagnosis\n"
                        "4. Current Treatments\n"
                        "5. Recommendations\n"
                        "6. Summary\n"
                    )
                )

                footer = (
                    record.report_template[0].footer if record.report_template else ""
                )
                header = (
                    record.report_template[0].header if record.report_template else ""
                )
                logo = record.report_template[0].logo if record.report_template else ""
                record.generated_report = self.get_information_from_chatgpt3(
                    record.transcription_data, record.record_locale, sections
                )

                pdf_data = VoiceRecordAiActions.text_to_pdf_formatted(
                    record.generated_report, footer=footer, header=header, logo=logo
                )
                docx_data = VoiceRecordAiActions.html_to_docx_formatted(
                    record.generated_report, footer=footer, header=header, logo=logo
                )

                formatted_date = datetime.datetime.now().strftime("%y%m%d")
                sequence_number = (
                        record.env["ir.sequence"].next_by_code(
                            "wrrrit.ai.voice_record.pdf.sequence"
                        )
                        or "0001"
                )
                pdf_filename = f"{record.name}-{formatted_date}-{sequence_number}.pdf"
                docx_filename = f"{record.name}-{formatted_date}-{sequence_number}.docx"

                record.write(
                    {
                        "generated_file_name": pdf_filename,
                        "generated_file": pdf_data,
                        "generated_docx_name": docx_filename,
                        "generated_docx": docx_data,
                    }
                )

                report_name = f"{record.name}-{formatted_date}-{sequence_number}-{record.record_locale}"
                record.env["wrrrit.ai.voice_record.pdf"].create(
                    {
                        "name": report_name,
                        "pdf_report_data": pdf_data,
                        "voice_record_id": record.id,
                    }
                )
            except Exception as e:
                _logger.error(
                    f"An error occurred during report generation for record ID {record.id}: {e}"
                )

    @staticmethod
    def html_to_docx_formatted(html_text, header, footer, logo):
        try:
            # Convert HTML to Markdown
            # md_text = markdown.markdown(html_text)
            md_text = html_text
            # Add Header and Footer to the Markdown text
            header = header if header and len(header) > 1 else ""
            md_text = md_text if md_text and len(md_text) > 1 else ""
            footer = footer if footer and len(footer) > 1 else ""

            md_text = f"# {header}\n\n{md_text}\n\n{footer}"

            # Convert Markdown to DOCX using pypandoc
            doc_path = f"document_{uuid.uuid4()}.docx"
            pypandoc.convert_text(md_text, "docx", format="html", outputfile=doc_path)

            if logo:
                # Generate a unique file name for the logo
                logo_filename = f"logo_{uuid.uuid4()}.png"
                logo_path = os.path.join(os.getcwd(), logo_filename)

                with open(logo_path, "wb") as logo_file:
                    logo_file.write(base64.b64decode(logo))

                # Open the DOCX file and add the logo
                doc = Document(doc_path)
                doc.paragraphs[
                    0
                ].clear()  # Clear the first paragraph which had the 'False' statement
                run = doc.paragraphs[0].add_run()
                run.add_picture(
                    logo_path, width=Inches(1.5)
                )  # Adjust width as necessary
                doc.save(doc_path)
                os.remove(logo_path)  # Delete the logo file if it was created

            # Optionally: Convert the DOCX file to base64 if needed
            with open(doc_path, "rb") as doc_file:
                doc_b64 = base64.b64encode(doc_file.read())

            # Optional: remove the doc file after converting it to base64
            os.remove(doc_path)

            return doc_b64

        except Exception as e:
            print(f"Failed to generate DOCX: {str(e)}")
            return None
